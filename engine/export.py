"""导出模块 — Markdown + Word。"""

from __future__ import annotations

import json
import re
from dataclasses import asdict, dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from .debate import GroupDiscussion, RoundResult, ROUNDS


def now_iso() -> str:
    return datetime.now(timezone.utc).astimezone().isoformat(timespec="seconds")


def export_markdown(
    topic: str,
    session_id: str,
    started_at: str,
    completed_at: str,
    winner_name: str,
    winner_stance: str,
    synthesis: str,
    groups: list[GroupDiscussion],
    review_responses: list[dict[str, str]],
    out_dir: Path,
) -> Path:
    """Export full debate results to Markdown."""
    md_path = out_dir / session_id / "debate.md"
    md_path.parent.mkdir(parents=True, exist_ok=True)

    lines = [
        f"# 圆桌派：{topic}",
        f"",
        f"- 开始时间：{started_at}",
        f"- 完成时间：{completed_at}",
        f"- 最终胜者：**{winner_name}**",
        f"",
    ]

    # Winner's original stance
    if winner_stance:
        lines.extend([
            f"## 胜者原始立场",
            f"",
            f"**{winner_name}**：",
            f"",
            winner_stance,
            f"",
        ])

    # Winner's synthesis
    if synthesis:
        lines.extend([
            f"## 思想结晶：{winner_name}的综合",
            f"",
            f"*（以下为{winner_name}在阅读了200位历史人物的回答后，对问题的综合判断）*",
            f"",
            synthesis,
            f"",
        ])

    # Group discussions
    for gi, group in enumerate(groups):
        agent_names = [a.name_zh for a in group.agents]
        lines.extend([
            f"## 第 {gi + 1} 组：{'、'.join(agent_names)}",
            f"",
        ])

        for rr in group.rounds:
            round_info = ROUNDS[rr.round_num - 1] if rr.round_num <= len(ROUNDS) else (rr.round_num, "?", "", 0, 0)
            lines.append(f"### 第 {rr.round_num} 轮（{round_info[1]}）")
            lines.append("")

            for utt in rr.utterances:
                change = " **[改票]**" if utt.changed else ""
                lines.append(f"**{utt.agent_name}**{change} → 认为「{utt.choice}」的回答最好：")
                lines.append(f"{utt.content}")
                lines.append("")

            if rr.vote_counts:
                lines.append("**票数统计**：")
                for name, count in rr.vote_counts.items():
                    marker = " ←" if count == max(rr.vote_counts.values()) else ""
                    lines.append(f"- {name}：{count}票{marker}")
                lines.append("")

        if group.final_votes:
            adv_name = next((a.name_zh for a in group.agents if a.id == group.advancers[0]), "?")
            lines.extend([
                f"**本组胜出：{adv_name}**",
                f"",
            ])

    # Global vote
    if review_responses:
        first = review_responses[0]
        if first.get("content", "").startswith("投给"):
            lines.extend([f"## 全民投票（200人）", f""])
            vote_counts: dict[str, int] = {}
            for resp in review_responses:
                content = resp["content"]
                match = re.search(r'投给「([^」]+)」', content)
                if match:
                    choice = match.group(1)
                    vote_counts[choice] = vote_counts.get(choice, 0) + 1
            lines.append("**最终票数**：")
            for name, count in sorted(vote_counts.items(), key=lambda x: -x[1]):
                marker = " ← 冠军" if count == max(vote_counts.values()) else ""
                lines.append(f"- {name}：{count}票{marker}")
            lines.append("")
            lines.append("**投票详情（精选）**：")
            for resp in review_responses[:40]:
                lines.append(f"- **{resp['agent_name']}**：{resp['content'][:200]}")
            lines.append("")

    md_path.write_text("\n".join(lines), encoding="utf-8")
    return md_path


def save_checkpoint(out_dir: Path, session_id: str, data: dict) -> Path:
    """Save checkpoint JSON."""
    cp_path = out_dir / session_id / "checkpoint.json"
    cp_path.parent.mkdir(parents=True, exist_ok=True)
    serializable = {}
    for k, v in data.items():
        if isinstance(v, list) and v and hasattr(v[0], '__dataclass_fields__'):
            serializable[k] = [asdict(item) for item in v]
        elif hasattr(v, '__dataclass_fields__'):
            serializable[k] = asdict(v)
        else:
            serializable[k] = v
    cp_path.write_text(json.dumps(serializable, ensure_ascii=False, indent=2))
    return cp_path


def export_word(
    topic: str,
    session_id: str,
    started_at: str,
    completed_at: str,
    winner_name: str,
    winner_stance: str,
    synthesis: str,
    groups: list[GroupDiscussion],
    review_responses: list[dict[str, str]],
    out_dir: Path,
) -> Path:
    """Export full debate results to Word document."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # 标题页
    doc.add_heading(f"圆桌派：{topic}", level=0)
    p = doc.add_paragraph()
    p.add_run(f"开始时间：{started_at}\n").font.size = Pt(10)
    p.add_run(f"完成时间：{completed_at}\n").font.size = Pt(10)
    p.add_run(f"最终胜者：{winner_name}").bold = True

    # 胜者立场
    if winner_stance:
        doc.add_heading("胜者原始立场", level=1)
        doc.add_paragraph(f"{winner_name}：", style='Intense Quote')
        doc.add_paragraph(winner_stance)

    # 思想综合
    if synthesis:
        doc.add_heading(f"思想结晶：{winner_name}的综合", level=1)
        doc.add_paragraph(
            f"（以下为{winner_name}在阅读了所有历史人物的回答后，对问题的综合判断）",
            style='Intense Quote',
        )
        doc.add_paragraph(synthesis)

    # 小组辩论
    for gi, group in enumerate(groups):
        agent_names = [a.name_zh for a in group.agents]
        doc.add_heading(f"第 {gi + 1} 组：{'、'.join(agent_names)}", level=1)

        for rr in group.rounds:
            round_info = ROUNDS[rr.round_num - 1] if rr.round_num <= len(ROUNDS) else (rr.round_num, "?", "", 0, 0)
            doc.add_heading(f"第 {rr.round_num} 轮（{round_info[1]}）", level=2)

            for utt in rr.utterances:
                change = " [改票]" if utt.changed else ""
                p = doc.add_paragraph()
                run = p.add_run(f"{utt.agent_name}{change}")
                run.bold = True
                p.add_run(f" → 认为「{utt.choice}」的回答最好：")
                doc.add_paragraph(utt.content)
                doc.add_paragraph()  # 空行

            if rr.vote_counts:
                doc.add_paragraph("票数统计：", style='List Bullet')
                for name, count in rr.vote_counts.items():
                    marker = " ←" if count == max(rr.vote_counts.values()) else ""
                    doc.add_paragraph(f"{name}：{count}票{marker}", style='List Bullet 2')

        if group.final_votes:
            adv_name = next((a.name_zh for a in group.agents if a.id == group.advancers[0]), "?")
            doc.add_paragraph(f"本组胜出：{adv_name}", style='Intense Quote')

    # 全民投票
    if review_responses:
        first = review_responses[0]
        if first.get("content", "").startswith("投给"):
            doc.add_heading("全民投票", level=1)
            vote_counts: dict[str, int] = {}
            for resp in review_responses:
                content = resp["content"]
                match = re.search(r'投给「([^」]+)」', content)
                if match:
                    choice = match.group(1)
                    vote_counts[choice] = vote_counts.get(choice, 0) + 1

            doc.add_paragraph("最终票数：")
            for name, count in sorted(vote_counts.items(), key=lambda x: -x[1]):
                marker = " ← 冠军" if count == max(vote_counts.values()) else ""
                doc.add_paragraph(f"{name}：{count}票{marker}", style='List Bullet')

            doc.add_paragraph("投票详情（精选）：")
            for resp in review_responses[:30]:
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(f"{resp['agent_name']}")
                run.bold = True
                p.add_run(f"：{resp['content'][:150]}")

    # 保存
    docx_path = out_dir / session_id / "debate.docx"
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    return docx_path
